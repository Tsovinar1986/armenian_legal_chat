# src/services/case_export.py
import os
from datetime import datetime
from typing import List, Dict, Optional

from docx import Document


class CaseExportService:
    """
    Service to export legal cases to text files and manage case documentation.
    """
    
    def __init__(self, export_dir: str = "exports"):
        self.export_dir = export_dir
        self._ensure_export_dir()
    
    def _ensure_export_dir(self):
        """Create export directory if it doesn't exist"""
        if not os.path.exists(self.export_dir):
            os.makedirs(self.export_dir)
            print(f"✅ Created export directory: {self.export_dir}")
    
    def export_similar_cases(self, cases: List[Dict], query: str) -> str:
        """
        Export similar cases to a text file.
        
        Args:
            cases: List of case dictionaries
            query: The original search query
            
        Returns:
            Path to the exported file
        """
        if not cases:
            print("⚠️ No cases to export")
            return None
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"similar_cases_{timestamp}.txt"
        filepath = os.path.join(self.export_dir, filename)
        
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write("=" * 80 + "\n")
                f.write("ՆՄԱՆ ԻՐԱՎԱԿԱՆ ԳՈՐԾԵՐԻ ՀԱՇՎԵՏՎՈՒԹՅՈՒՆ\n")
                f.write("SIMILAR LEGAL CASES REPORT\n")
                f.write("=" * 80 + "\n\n")
                
                f.write(f"📅 Հաշվետվության ամսաթիվ / Report Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"🔍 Որոնման հարց / Search Query: {query}\n")
                f.write(f"📊 Գտնված գործերի քանակ / Total Cases Found: {len(cases)}\n\n")
                
                f.write("=" * 80 + "\n")
                f.write("ԳՈՐԾԵՐԻ ՄԱՆՐԱՄԱՍՆԵՐ / CASE DETAILS\n")
                f.write("=" * 80 + "\n\n")
                
                for idx, case in enumerate(cases, 1):
                    f.write(f"\n{'─' * 80}\n")
                    f.write(f"ԳՈՐԾ #{idx} / CASE #{idx}\n")
                    f.write(f"{'─' * 80}\n\n")
                    
                    # Case number
                    case_num = case.get('unique_number') or case.get('case_number') or 'N/A'
                    f.write(f"📌 Գործի համարը / Case Number: {case_num}\n")
                    
                    # Lawyer name
                    lawyer = case.get('lawyer_name') or case.get('lawyer') or 'Նշված չէ / Not specified'
                    if lawyer == "(NULL)":
                        lawyer = "Նշված չէ / Not specified"
                    f.write(f"⚖️ Փաստաբանի անուն / Lawyer Name: {lawyer}\n")
                    
                    # Classification
                    classification = case.get('civil_case_classifier') or case.get('classification') or 'N/A'
                    f.write(f"📋 Դասակարգում / Classification: {classification}\n")
                    
                    # Link
                    link = case.get('link') or 'N/A'
                    f.write(f"🌐 Հղում / Link: {link}\n")
                    
                    # Verdict (if available)
                    verdict = case.get('verdict_text') or case.get('Verdict_Text') or ''
                    if verdict:
                        f.write(f"\n📄 Վճիռ / Verdict:\n")
                        f.write(f"{verdict}\n")
                    
                    # Judicial prehistory
                    prehistory = case.get('judicial_prehistory') or case.get('prehistory') or ''
                    if prehistory:
                        f.write(f"\n📖 Գործի պատմություն / Judicial Prehistory:\n")
                        f.write(f"{prehistory}\n")
                    
                    # Status (if available)
                    status = case.get('status') or case.get('outcome') or ''
                    if status:
                        f.write(f"\n✅ Վճիռի ելք / Case Outcome: {status}\n")
                    
                    f.write("\n")
            
            print(f"✅ Cases exported to: {filepath}")
            return filepath
            
        except Exception as e:
            print(f"❌ Error exporting cases: {e}")
            return None
    
    def export_approved_cases(self, approved_cases: List[Dict]) -> str:
        """
        Export approved/successful cases to a text file.
        
        Args:
            approved_cases: List of approved case dictionaries
            
        Returns:
            Path to the exported file
        """
        if not approved_cases:
            print("⚠️ No approved cases to export")
            return None
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"approved_cases_{timestamp}.txt"
        filepath = os.path.join(self.export_dir, filename)
        
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write("=" * 80 + "\n")
                f.write("ՀԱՋՈՂՎԱԾ ԻՐԱՎԱԿԱՆ ԳՈՐԾԵՐԻ ՀԱՇՎԵՏՎՈՒԹՅՈՒՆ\n")
                f.write("APPROVED/SUCCESSFUL LEGAL CASES REPORT\n")
                f.write("=" * 80 + "\n\n")
                
                f.write(f"📅 Հաշվետվության ամսաթիվ / Report Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"📊 Հաջողված գործերի քանակ / Total Approved Cases: {len(approved_cases)}\n\n")
                
                # Statistics
                lawyers_count = len(set(c.get('lawyer_name', 'N/A') for c in approved_cases))
                f.write(f"👨‍⚖️ Փաստաբանների քանակ / Number of Lawyers: {lawyers_count}\n\n")
                
                f.write("=" * 80 + "\n")
                f.write("ԳՈՐԾԵՐԻ ԸՆԴԱՄԵՆԸ / CASES SUMMARY\n")
                f.write("=" * 80 + "\n\n")
                
                for idx, case in enumerate(approved_cases, 1):
                    f.write(f"\n{'─' * 80}\n")
                    f.write(f"ՀԱՋՈՂՎԱԾ ԳՈՐԾ #{idx} / APPROVED CASE #{idx}\n")
                    f.write(f"{'─' * 80}\n\n")
                    
                    # Case number
                    case_num = case.get('unique_number') or case.get('case_number') or 'N/A'
                    f.write(f"📌 Գործի համարը / Case Number: {case_num}\n")
                    
                    # Lawyer name (emphasized)
                    lawyer = case.get('lawyer_name') or case.get('lawyer') or 'Նշված չէ'
                    if lawyer == "(NULL)":
                        lawyer = "Նշված չէ"
                    f.write(f"⭐ Փաստաբանի անուն / Lawyer Name: {lawyer}\n")
                    f.write(f"📞 Փաստաբանի բաժանմունք / Lawyer Department: {case.get('department', 'N/A')}\n")
                    
                    # Classification
                    classification = case.get('civil_case_classifier') or case.get('classification') or 'N/A'
                    f.write(f"📋 Դասակարգում / Classification: {classification}\n")
                    
                    # Approval/Success indicator
                    approval_status = case.get('approval_status') or case.get('status') or 'Հաստատված / Approved'
                    f.write(f"✅ Վճիռի տեսակ / Verdict Type: {approval_status}\n")
                    
                    # Link
                    link = case.get('link') or 'N/A'
                    f.write(f"🌐 Դատական հղում / Case Link: {link}\n")
                    
                    # Description
                    description = case.get('judicial_prehistory') or case.get('description') or ''
                    if description:
                        # Limit to first 500 chars
                        if len(description) > 500:
                            description = description[:500] + "..."
                        f.write(f"\n📝 Գործի նկարագրություն / Description:\n")
                        f.write(f"{description}\n")
                    
                    f.write("\n")
            
            print(f"✅ Approved cases exported to: {filepath}")
            return filepath
            
        except Exception as e:
            print(f"❌ Error exporting approved cases: {e}")
            return None
    
    def export_similar_cases_docx(self, cases: List[Dict], query: str) -> str:
        """
        Export similar cases (same case type) to a Word document with links.

        Args:
            cases: List of case dictionaries
            query: The original search query

        Returns:
            Path to the exported .docx file
        """
        if not cases:
            print("⚠️ No cases to export")
            return None

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"similar_cases_{timestamp}.docx"
        filepath = os.path.join(self.export_dir, filename)

        try:
            doc = Document()
            doc.add_heading("ՆՄԱՆ ԻՐԱՎԱԿԱՆ ԳՈՐԾԵՐԻ ՀԱՇՎԵՏՎՈՒԹՅՈՒՆ / Similar Legal Cases Report", level=1)

            meta = doc.add_paragraph()
            meta.add_run(f"Report Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n").bold = False
            meta.add_run(f"Search Query: {query}\n")
            meta.add_run(f"Total Cases Found: {len(cases)}")

            for idx, case in enumerate(cases, 1):
                doc.add_heading(f"Case #{idx}", level=2)

                case_num = case.get('unique_number') or case.get('case_number') or 'N/A'
                doc.add_paragraph(f"Case Number: {case_num}")

                lawyer = case.get('lawyer_name') or case.get('lawyer') or 'Not specified'
                if lawyer == "(NULL)":
                    lawyer = "Not specified"
                doc.add_paragraph(f"Lawyer Name: {lawyer}")

                classification = case.get('civil_case_classifier') or case.get('classification') or 'N/A'
                doc.add_paragraph(f"Classification / Case Type: {classification}")

                link = case.get('link') or 'N/A'
                doc.add_paragraph(f"Link: {link}")

                prehistory = case.get('judicial_prehistory') or case.get('prehistory') or ''
                if prehistory:
                    doc.add_paragraph("Judicial Prehistory:")
                    doc.add_paragraph(prehistory)

            doc.save(filepath)
            print(f"✅ Cases exported to: {filepath}")
            return filepath

        except Exception as e:
            print(f"❌ Error exporting cases to docx: {e}")
            return None

    def export_approved_cases_docx(self, approved_cases: List[Dict], top_lawyers: Optional[List] = None) -> str:
        """
        Export approved/successful cases, and the ranking of lawyers by approved
        case count, to a Word document.

        Args:
            approved_cases: List of approved case dictionaries
            top_lawyers: Optional list of (lawyer_name, stats) tuples, stats containing
                at least a 'count' key, sorted by most approved cases first. Agency
                name and location are not yet tracked in the source data, so only the
                lawyer name and case count are included.

        Returns:
            Path to the exported .docx file
        """
        if not approved_cases:
            print("⚠️ No approved cases to export")
            return None

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"approved_cases_{timestamp}.docx"
        filepath = os.path.join(self.export_dir, filename)

        try:
            doc = Document()
            doc.add_heading("ՀԱՋՈՂՎԱԾ ԻՐԱՎԱԿԱՆ ԳՈՐԾԵՐ / Approved Legal Cases Report", level=1)

            meta = doc.add_paragraph()
            meta.add_run(f"Report Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            meta.add_run(f"Total Approved Cases: {len(approved_cases)}")

            if top_lawyers:
                doc.add_heading("Top Lawyer by Approved Cases", level=2)
                top_name, top_stats = top_lawyers[0]
                doc.add_paragraph(
                    f"{top_name} — {top_stats['count']} approved case(s) "
                    f"(agency name and location not available in current dataset)"
                )

                if len(top_lawyers) > 1:
                    doc.add_heading("Other Lawyers (Ranked)", level=2)
                    for name, stats in top_lawyers[1:]:
                        doc.add_paragraph(f"{name} — {stats['count']} approved case(s)", style="List Bullet")

            doc.add_heading("Case Details", level=2)
            for idx, case in enumerate(approved_cases, 1):
                doc.add_heading(f"Approved Case #{idx}", level=3)

                case_num = case.get('unique_number') or case.get('case_number') or 'N/A'
                doc.add_paragraph(f"Case Number: {case_num}")

                lawyer = case.get('lawyer_name') or case.get('lawyer') or 'Not specified'
                if lawyer == "(NULL)":
                    lawyer = "Not specified"
                doc.add_paragraph(f"Lawyer Name: {lawyer}")

                classification = case.get('civil_case_classifier') or case.get('classification') or 'N/A'
                doc.add_paragraph(f"Classification / Case Type: {classification}")

                link = case.get('link') or 'N/A'
                doc.add_paragraph(f"Link: {link}")

            doc.save(filepath)
            print(f"✅ Approved cases exported to: {filepath}")
            return filepath

        except Exception as e:
            print(f"❌ Error exporting approved cases to docx: {e}")
            return None

    def get_export_directory(self) -> str:
        """Get the export directory path"""
        return os.path.abspath(self.export_dir)
    
    def list_exports(self) -> List[str]:
        """List all exported files"""
        try:
            files = os.listdir(self.export_dir)
            exported_files = [f for f in files if f.endswith('.txt') or f.endswith('.docx')]
            return sorted(exported_files, reverse=True)
        except Exception as e:
            print(f"⚠️ Error listing exports: {e}")
            return []
